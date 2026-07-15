def parse_docx(file_path: str) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text_parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            text_parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            rt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if rt:
                text_parts.append(rt)
    if not text_parts:
        raise ValueError("No text could be extracted from the DOCX file")
    return "\n".join(text_parts)


def parse_pdf(file_path: str) -> str:
    import pdfplumber
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    text_parts.append(t)
    except Exception as e:
        raise ValueError(f"Failed to read PDF file: {e}") from e
    if not text_parts:
        raise ValueError("No text could be extracted from the PDF file")
    return "\n".join(text_parts)


def parse_image(file_path: str) -> str:
    import base64
    import os
    from langchain_core.messages import HumanMessage
    from app.agents.llm_client import get_chat_model
    from app.core.prompts import OCR_PROMPT

    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    mime = "jpeg" if ext in ("jpg", "jpeg") else ext

    with open(file_path, "rb") as f:
        image_b64 = base64.b64encode(f.read()).decode("utf-8")

    message = HumanMessage(content=[
        {"type": "text", "text": OCR_PROMPT},
        {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{image_b64}"}},
    ])
    text = get_chat_model().invoke([message]).content.strip()
    if not text:
        raise ValueError("No text could be extracted from the image")
    return text


def parse_document(file_path: str, file_ext: str) -> str:
    ext = file_ext.lower()
    if ext in (".docx", ".doc"):
        return parse_docx(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in (".png", ".jpg", ".jpeg"):
        return parse_image(file_path)
    raise ValueError(f"Unsupported file format: {ext}")
